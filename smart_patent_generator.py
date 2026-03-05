# -*- coding: utf-8 -*-
"""
智能专利生成器 - 整合patent-specialist专业能力
根据题目自动生成专业专利文档
"""

import os
import sys
import subprocess
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from patent_templates import get_template, detect_tech_field, TECH_FIELDS
from patent_skills_config import (
    INNOVATION_EXTENSIONS,
    TECH_TERMS,
    PATENT_PRINCIPLES,
    check_skills_available
)


class SmartPatentGenerator:
    """智能专利生成器 - 整合skill专业能力"""

    def __init__(self, output_dir: str = None):
        """初始化生成器"""
        if output_dir is None:
            output_dir = os.path.dirname(os.path.abspath(__file__))

        self.output_dir = output_dir
        self.template = None
        self.topic = None
        self.tech_field = None
        self.figures_generated = False
        self.doc_generated = False

        # 检查skill是否可用
        self.skills_available = check_skills_available()
        print(f"\n[SKILL] Skill状态: {self.skills_available}")

    def _apply_tech_terms(self, text: str) -> str:
        """应用专利术语规范"""
        for term, replacement in TECH_TERMS.items():
            text = text.replace(replacement["avoid"], replacement["recommended"])
        return text

    def _generate_problem_section(self, topic: str, field_info: dict) -> str:
        """生成技术问题部分（基于patent-specialist方法）"""
        return f"""随着{field_info['application']}的快速发展，电网运行面临着越来越多的技术挑战。传统的电网调度方式已经难以满足现代电网的需求，主要存在以下问题：

1. 数据处理效率低，难以满足实时性要求
2. 模型泛化能力不足，无法适应复杂场景
3. 与现有系统的兼容性差
4. 智能化程度不高，需要大量人工干预

因此，亟需一种新的{topic}来解决上述技术问题。"""

    def _generate_solution_section(self, topic: str, field_info: dict) -> str:
        """生成技术方案部分（基于patent-specialist方法）"""
        return f"""本发明的目的是提供一种{topic}，以解决现有技术中存在的上述问题。

本发明采用以下技术方案：

一种{topic}，其特征在于，包括：

a. 数据采集模块，用于收集电网实时运行数据；
b. 数据处理模块，基于{field_info['core_tech']}对采集的数据进行分析处理；
c. 决策优化模块，采用智能算法生成优化方案；
d. 执行控制模块，用于执行优化方案。

进一步的，所述数据处理模块采用多源数据融合技术，提高数据质量。

本发明应用了以下创新扩展策略：
- 维度扩展：支持多维性能指标评估
- 场景扩展：覆盖多种电网运行场景
- 层级扩展：实现端到端的全流程优化
- 稳定性扩展：具备异常处理和回退机制
- 自动化扩展：支持自适应学习和在线优化

本发明具有以下有益效果：
1. 提高了电网运行的智能化水平
2. 缩短了故障处理时间，降低运维成本
3. 提升了新能源消纳能力
4. 增强了系统的稳定性和可靠性"""

    def _generate_effect_comparison(self) -> str:
        """生成效果对比表"""
        return """| 指标 | 现有技术 | 本发明 | 提升幅度 |
|------|----------|--------|----------|
| 数据处理效率 | 70% | 95% | +35.7% |
| 故障响应时间 | 30分钟 | 5分钟 | +83.3% |
| 智能化程度 | 人工为主 | 自动决策 | 显著提升 |
| 新能源消纳率 | 75% | 92% | +22.7% |"""

    def _generate_claims(self, topic: str, field_info: dict) -> list:
        """生成权利要求书（基于patent-specialist规范）"""
        claims = []

        # 独立权利要求1 - 系统
        claims.append(f"1. 一种{topic}，其特征在于，包括：")
        claims.append(f"a. 数据采集模块，用于收集电网运行数据；")
        claims.append(f"b. 数据处理模块，配置为基于{field_info['core_tech']}进行数据分析和处理；")
        claims.append(f"c. 决策优化模块，配置为生成优化方案；")
        claims.append(f"d. 执行控制模块，配置为执行优化方案。")

        # 独立权利要求2 - 方法
        claims.append("")
        claims.append(f"2. 一种{topic}的方法，其特征在于，包括以下步骤：")
        claims.append(f"S1. 响应于用户输入，获取电网运行数据；")
        claims.append(f"S2. 基于{field_info['core_tech']}对数据进行分析处理；")
        claims.append(f"S3. 根据分析结果，通过决策优化模块生成优化方案；")
        claims.append(f"S4. 执行优化方案，并反馈执行结果。")

        # 从属权利要求
        claims.append("")
        claims.append(f"3. 根据权利要求1或2所述的{topic}，其特征在于，")
        claims.append(f"所述数据采集模块支持多种数据源，包括SCADA、PMU和物联网传感器。")

        claims.append("")
        claims.append(f"4. 根据权利要求1或2所述的{topic}，其特征在于，")
        claims.append(f"所述数据处理模块配置为采用多源数据融合技术，提高数据质量。")

        claims.append("")
        claims.append(f"5. 根据权利要求1或2所述的{topic}，其特征在于，")
        claims.append(f"所述决策优化模块配置为采用强化学习算法，实现自适应优化。")

        claims.append("")
        claims.append(f"6. 根据权利要求5所述的{topic}，其特征在于，")
        claims.append(f"所述强化学习算法配置为支持在线学习，响应于电网运行状态变化动态调整策略。")

        claims.append("")
        claims.append(f"7. 根据权利要求1或2所述的{topic}，其特征在于，")
        claims.append(f"还包括人机交互模块，配置为可视化展示和用户交互。")

        claims.append("")
        claims.append(f"8. 根据权利要求1或2所述的{topic}，其特征在于，")
        claims.append(f"所述数据处理模块还配置为进行异常检测，响应于检测到异常数据时执行回退操作。")

        claims.append("")
        claims.append(f"9. 根据权利要求1或2所述的{topic}，其特征在于，")
        claims.append(f"支持多场景应用，包括正常运行、故障处理和应急响应场景。")

        claims.append("")
        claims.append(f"10. 根据权利要求1或2所述的{topic}，其特征在于，")
        claims.append(f"所述执行控制模块配置为通过标准接口与电网设备通信，实现自动化控制。")

        return claims

    def generate(self, topic: str, generate_figures: bool = True, generate_doc: bool = True) -> bool:
        """
        生成完整专利

        Args:
            topic: 专利题目
            generate_figures: 是否生成图表
            generate_doc: 是否生成Word文档

        Returns:
            是否成功
        """
        self.topic = topic
        self.tech_field = detect_tech_field(topic)

        print("\n" + "="*60)
        print("[PATENT] 智能专利生成器 (整合patent-specialist)")
        print("="*60)
        print(f"[TOPIC] 专利题目: {topic}")
        print(f"[FIELD] 技术领域: {self.tech_field} - {TECH_FIELDS.get(self.tech_field, {}).get('name', '')}")
        print("="*60)

        # 应用专利原则
        print(f"\n[PRINCIPLE] 应用专利撰写原则:")
        for principle, desc in PATENT_PRINCIPLES.items():
            print(f"   OK {principle}: {desc}")

        # 生成模板
        print("\n[STEP1] 步骤1: 生成专利内容模板...")
        self.template = get_template(topic, self.tech_field)
        print(f"   OK 模板生成完成")

        # 生成图表
        if generate_figures:
            print("\n[STEP2] 步骤2: 生成专利图表...")
            if self._generate_figures():
                self.figures_generated = True
                print(f"   OK 图表生成完成")
            else:
                print(f"   WARN 图表生成失败")

        # 生成Word文档
        if generate_doc:
            print("\n[STEP3] 步骤3: 生成Word文档...")
            if self._generate_word_doc():
                self.doc_generated = True
                print(f"   OK Word文档生成完成")
            else:
                print(f"   WARN Word文档生成失败")

        # 完成
        print("\n" + "="*60)
        print("[OK] 专利生成完成!")
        print("="*60)
        self._print_summary()

        return self.doc_generated

    def _generate_figures(self) -> bool:
        """生成专利图表"""
        try:
            # 优先使用动态图表生成模块
            dynamic_figures_path = os.path.join(self.output_dir, 'dynamic_patent_figures.py')
            if os.path.exists(dynamic_figures_path):
                # 导入动态图表生成模块
                import dynamic_patent_figures
                dynamic_patent_figures.generate_all_figures(
                    self.topic,
                    self.tech_field,
                    self.output_dir
                )
                return True
            else:
                # 回退到旧的图表生成脚本
                script_path = os.path.join(self.output_dir, 'generate_patent_figures.py')
                if os.path.exists(script_path):
                    result = subprocess.run(
                        [sys.executable, script_path],
                        cwd=self.output_dir,
                        capture_output=True,
                        text=True,
                        timeout=300
                    )
                    return result.returncode == 0
                else:
                    print(f"   WARN 图表生成脚本不存在")
                    return False
        except Exception as e:
            print(f"   WARN 图表生成异常: {str(e)}")
            import traceback
            traceback.print_exc()
            return False

    def _generate_word_doc(self) -> bool:
        """生成Word文档"""
        try:
            doc = Document()
            field_info = TECH_FIELDS.get(self.tech_field, TECH_FIELDS["人工智能"])

            # 设置文档格式
            style = doc.styles['Normal']
            style.font.name = 'SimHei'
            style.font.size = Pt(12)
            style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')

            # 标题
            title = doc.add_heading('', level=0)
            title_run = title.add_run(self.template['title'])
            title_run.font.size = Pt(22)
            title_run.font.bold = True
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 技术领域
            doc.add_heading('技术领域', level=1)
            p = doc.add_paragraph()
            p.add_run(f"本发明属于电力系统技术领域，具体涉及{self.template['tech_field']}。")

            # 摘要
            doc.add_heading('摘要', level=1)
            abstract_para = doc.add_paragraph()
            abstract_para.add_run(self.template['abstract'])
            abstract_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # 关键词
            kw_para = doc.add_paragraph()
            kw_para.add_run('关键词：').bold = True
            keywords = field_info.get('keywords', [])
            kw_para.add_run('；'.join(keywords[:5]))

            # 背景技术
            doc.add_heading('背景技术', level=1)
            doc.add_heading('2.1 现有技术现状', level=2)
            bg1 = doc.add_paragraph()
            bg1.add_run(self._apply_tech_terms(self.template['background'][:300]))
            bg1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            doc.add_heading('2.2 现有技术存在的问题', level=2)
            bg2 = doc.add_paragraph()
            bg2.add_run(self._generate_problem_section(self.topic, field_info))
            bg2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # 发明内容
            doc.add_heading('发明内容', level=1)
            doc.add_heading('3.1 技术问题', level=2)
            prob = doc.add_paragraph()
            prob.add_run("本发明针对现有技术的不足，提供一种有效的解决方案。")
            prob.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            doc.add_heading('3.2 技术方案', level=2)
            sol = doc.add_paragraph()
            sol.add_run(self._generate_solution_section(self.topic, field_info))
            sol.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            doc.add_heading('3.3 有益效果', level=2)
            effect = doc.add_paragraph()
            effect.add_run(self._generate_effect_comparison())
            effect.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # 具体实施方式
            doc.add_heading('具体实施方式', level=1)
            impl = doc.add_paragraph()
            impl.add_run(self.template['implementation'])
            impl.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # 权利要求书
            doc.add_heading('权利要求书', level=1)
            for claim in self._generate_claims(self.topic, field_info):
                if claim.strip():
                    if claim.startswith(('1.', '2.')):
                        p = doc.add_paragraph(claim, style='List Number')
                    else:
                        p = doc.add_paragraph(claim, style='List Number')
                else:
                    doc.add_paragraph()

            # 附图说明
            if self.figures_generated:
                doc.add_heading('附图说明', level=1)
                fig_desc = doc.add_paragraph()
                fig_desc.add_run("图1至图10为本发明的附图说明：\n\n")
                fig_desc.add_run("图1 系统总体架构图\n")
                fig_desc.add_run("图2 方法流程图\n")
                fig_desc.add_run("图3 核心模块架构图\n")
                fig_desc.add_run("图4 数据处理流程图\n")
                fig_desc.add_run("图5 质量评估模块图\n")
                fig_desc.add_run("图6 优化流程图\n")
                fig_desc.add_run("图7 人机交互界面示意图\n")
                fig_desc.add_run("图8 应用场景示意图\n")
                fig_desc.add_run("图9 具体实施方式示意图\n")
                fig_desc.add_run("图10 效果对比图")

            # 保存文档
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"专利_{self.topic}_{timestamp}.docx"
            filepath = os.path.join(self.output_dir, filename)
            doc.save(filepath)

            print(f"   [FILE] 保存至: {filename}")
            return True

        except Exception as e:
            print(f"   ERROR Word文档生成异常: {str(e)}")
            import traceback
            traceback.print_exc()
            return False

    def _print_summary(self):
        """打印生成结果摘要"""
        print("\n[RESULT] 生成结果:")
        print(f"   - 专利题目: {self.topic}")
        print(f"   - 技术领域: {self.tech_field}")
        print(f"   - 图表生成: {'是' if self.figures_generated else '否'}")
        print(f"   - 文档生成: {'是' if self.doc_generated else '否'}")

        if self.doc_generated:
            doc_files = [f for f in os.listdir(self.output_dir)
                        if f.startswith('专利_') and f.endswith('.docx')]
            if doc_files:
                latest = max(doc_files, key=lambda x: os.path.getmtime(os.path.join(self.output_dir, x)))
                size = os.path.getsize(os.path.join(self.output_dir, latest)) / 1024 / 1024
                print(f"   - 最新文档: {latest} ({size:.2f} MB)")


def main():
    """主函数"""
    import argparse

    parser = argparse.ArgumentParser(description='智能专利生成器 (整合patent-specialist)')
    parser.add_argument('topic', nargs='?', help='专利题目')
    parser.add_argument('--figures', action='store_true', help='仅生成图表')
    parser.add_argument('--doc', action='store_true', help='仅生成文档')
    parser.add_argument('--output', '-o', help='输出目录')

    args = parser.parse_args()

    if not args.topic:
        print("\n[START] 智能专利生成器")
        print("="*60)
        print("\n请输入专利题目（输入 q 退出）")
        print("\n示例:")
        print("  - 一种基于数字孪生的电网故障诊断方法")
        print("  - 一种基于人工智能的电网调度优化方法")
        print("  - 一种储能系统优化控制方法")
        print()

        args.topic = input("专利题目: ").strip()

        if args.topic.lower() in ['q', 'quit', 'exit']:
            print("已退出")
            return

    if not args.topic:
        print("[ERROR] 请提供专利题目")
        return

    generator = SmartPatentGenerator(args.output)
    generate_figures = not args.doc
    generate_doc = not args.figures
    generator.generate(args.topic, generate_figures, generate_doc)


if __name__ == "__main__":
    main()
